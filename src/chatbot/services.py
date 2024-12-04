import json
import logging

# flake8:noqa
import os
from datetime import datetime
from io import BytesIO
from xml.dom.minidom import Document

import firebase_admin
import requests
import speech_recognition as sr
from docx import Document
from dotenv import load_dotenv
from fastapi.concurrency import run_in_threadpool
from fastapi.responses import FileResponse
from firebase_admin import credentials, storage
from linebot import LineBotApi
from linebot.models import TextSendMessage
from pydub import AudioSegment

import src.chatbot.models as models
from src.anythingllm.route import (
    creat_new_thread,
    create_workspace,
    get_documents,
    send_chat_message,
    update_embedding,
    upload_document,
)
from src.chatbot.database import (
    get_group_chat_records_by_id,
    save_group_chat_records_to_db,
)
from src.chatbot.message_templates.buttom_template import create_check_button
from src.chatbot.message_templates.carousel_template import (  # noqa
    create_houses_carousel,
)

from .logconfig import setup_logging

setup_logging()
load_dotenv()

base_url = os.getenv("BASE_URL")
api_key = os.getenv("API_KEY")
default_path = os.getenv("DEFAULT_PATH")
line_bot_api = LineBotApi(os.getenv("CHANNEL_ACCESS_TOKEN"))


def normal_chat(user_id, user_message):
    workspace = create_workspace(api_key, user_id)
    slug = workspace.get_slug()
    thread_id = creat_new_thread(slug)
    response = send_chat_message(api_key, slug, thread_id, user_message, mode="chat")
    return response


def save_group_chat_records(user_id, group_id, msg):
    profile = line_bot_api.get_group_member_profile(group_id, user_id)
    user_name = profile.display_name
    msg_detail = models.MsgDetail(
        UserId=user_id, UserName=user_name, MsgText=msg, Time=datetime.now()
    )
    try:
        result = save_group_chat_records_to_db(group_id, msg_detail)
    except Exception as e:
        result = f"Exception: {e}"
    return result


def group_chat_records_to_file(group_id):
    try:
        chat_record_list = get_group_chat_records_by_id(group_id)
        os.makedirs("./src/chatbot/group_chat_records/", exist_ok=True)
        with open(
            f"./src/chatbot/group_chat_records/{group_id}.json", "w"
        ) as file:  # noqa
            json.dump(chat_record_list, file, indent=4)
    except Exception as e:
        print(f"Failed to transfer chat records to file: {e}")


def generate_summarized_checklist(group_id):
    # Integrate anythingllm api to generate consensus doc
    workspace = create_workspace(api_key, workspace_name=group_id)
    upload_document(api_key=api_key, file_path=f"{default_path}/{group_id}.json")
    doc = get_documents()
    update_embedding(api_key, workspace.get_slug(), doc.get_doc_id())
    thread_id = creat_new_thread(workspace.get_slug())
    file_path = "./llm-promting/promt.txt"
    with open(file_path, "r", encoding="utf-8") as file:
        prompt = file.read()
        ans = send_chat_message(
            api_key,
            workspace.get_slug(),
            thread_id=thread_id,
            message=prompt,
            mode="chat",
        )
    return ans


# def summary_checklist():
#     buttom_message = create_check_button()
#     return buttom_message


def call_llm_api():
    result = "call_llm_api()"
    return result


# Recommendation related services
def house_recommendation():
    your_ip = os.getenv("HOUSE_RECOMMEND_API")
    user_id = 2
    url = f"{your_ip}/get_recommendation/0/{user_id}"
    try:
        response = requests.get(url)
        data_list = response.json()
        carousel_message = create_houses_carousel(data_list)
        result = carousel_message
    except requests.HTTPError as http_err:
        error_message = f"API Request Failed: {http_err}"
        result = error_message
    return result


async def voice_recognition(msg_id):
    result_content = ""
    try:
        message_content = line_bot_api.get_message_content(msg_id)
        m4a_audio_bytes_io = BytesIO()
        for chunk in message_content.iter_content():
            m4a_audio_bytes_io.write(chunk)
        m4a_audio_bytes_io.seek(0)

        request_audio = AudioSegment.from_file(m4a_audio_bytes_io, format="M4A")
        wav_audio_bytes_io = request_audio.export(format="wav")
        r = sr.Recognizer()
        with sr.AudioFile(wav_audio_bytes_io) as source:
            recognizer_audio = r.record(source)
        result_content = await run_in_threadpool(
            lambda: r.recognize_google(recognizer_audio, language="zh-TW")
        )
    except Exception as e:
        result_content = f"Failed to recognize audio. {e}"

    return result_content


async def handle_async_audio(event):
    try:
        msg_id = event.message.id
        recognized_text = await voice_recognition(msg_id)

        line_bot_api.reply_message(
            event.reply_token, TextSendMessage(text=recognized_text)
        )

    except Exception as e:
        print(f"Errors occurred when handling async audio: {e}")


def __save_document_to_firebase(file_path, group_id):
    cred = credentials.Certificate(
        "src/firebase/fang5-group-firebase-adminsdk-d81ae-4c09d6c55f.json"
    )
    firebase_admin.initialize_app(cred, {"storageBucket": "fang5-group.appspot.com"})

    blob_name = f"uploads/consensus-{group_id}.docx"  # Firebase Storage 中的檔案路徑

    # 獲取儲存桶（bucket）並上傳文件
    bucket = storage.bucket()
    blob = bucket.blob(blob_name)
    blob.upload_from_filename(file_path)

    # 設置該文件的公開訪問權限
    blob.make_public()

    # 獲取文件的公開 URL
    url = blob.public_url

    print(f"File uploaded successfully. URL: {url}")
    return url


def build_consensus_document(input_text, group_id):
    # Generate consensus document
    doc = Document()
    doc.add_heading("青銀共居協議書", 0).alignment = 1

    sections = input_text.strip().split("#")[1:]
    for section in sections:
        if section.strip():
            heading, *content = section.split("\n")
            doc.add_heading(heading.strip(), level=2)
            for line in content:
                if line.strip().startswith("-"):
                    line_content = line.strip().lstrip("-").strip()
                    doc.add_paragraph(line_content, style="ListBullet")
                elif line.strip().startswith("+"):
                    line_content = line.strip().lstrip("+").strip()
                    doc.add_paragraph(line_content, style="ListBullet")
                elif line.strip().startswith("*"):
                    line_content = line.strip().lstrip("*").strip()
                    doc.add_paragraph(line_content, style="ListBullet")
                else:
                    line_content = line.strip()
                    doc.add_paragraph(line_content)

    file_path = f"src/firebase/consensus-{group_id}.docx"
    doc.save(file_path)
    url = __save_document_to_firebase(file_path, group_id)
    return url
